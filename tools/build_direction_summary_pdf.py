from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/book/Documents/auto-cs-bot-ver-1")
OUT = ROOT / "deliverables" / "hk-sme-dm-sales-cockpit-direction-summary.docx"

FONT = "STHeiti"
TITLE_COLOR = RGBColor(16, 35, 51)
BLUE = RGBColor(34, 94, 138)
MUTED = RGBColor(92, 103, 112)
GREEN = RGBColor(37, 105, 82)
AMBER = RGBColor(125, 88, 20)
RED = RGBColor(142, 45, 45)
LIGHT_BLUE = "EAF3F8"
LIGHT_GREEN = "EAF5EF"
LIGHT_AMBER = "FFF4DF"
LIGHT_GRAY = "F4F6F8"


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in [
        ("Title", 24, TITLE_COLOR, 0, 8),
        ("Heading 1", 16, BLUE, 16, 7),
        ("Heading 2", 12.5, BLUE, 10, 5),
        ("Heading 3", 11.5, TITLE_COLOR, 8, 3),
    ]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.12


def add_para(doc, text="", size=10.5, bold=False, color=None, align=None, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, bold=True)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.12
        run = p.add_run(item)
        set_run_font(run, size=10.3)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    set_run_font(run, size=9.3, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if fill:
        shade_cell(cell, fill)


def add_comparison_table(doc):
    headers = ["項目", "Omnichat", "我們應該做"]
    rows = [
        ("定位", "大型品牌 chat commerce / omnichannel platform", "香港小店 DM 漏客偵測 + 成交跟進系統"),
        ("客戶", "Fortress、Watsons、FILA、Timberland 等大品牌", "每天 20-100 個 IG/WhatsApp 查詢的小店"),
        ("價格", "公開頁只寫 Quote，annual subscription，另計 API/message fee", "透明月費、可月付、低 setup 成本"),
        ("功能重點", "CDP、broadcast、journey、coupon/game、OMO sales", "inbox、AI 草稿、lead status、跟進提醒、每日報告"),
        ("風險", "功能完整但重，需要 onboarding 和整合", "要避免做成普通 chatbot，要先打中一個尖痛點"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    widths = [1.15, 2.7, 2.9]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, fill=LIGHT_GRAY)
        table.rows[0].cells[idx].width = Inches(widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, bold=(idx == 0))
            cells[idx].width = Inches(widths[idx])
    doc.add_paragraph()


def add_roadmap_table(doc):
    headers = ["階段", "要做什麼", "目的"]
    rows = [
        ("MVP 1", "DM Inbox Lite + AI 草稿 + 安全檢查", "先證明可以幫員工快回、不亂承諾"),
        ("MVP 2", "Lead 狀態：新查詢 / 高意向 / 等付款 / 已成交 / 投訴風險", "把 DM 變成可管理的 pipeline"),
        ("MVP 3", "每日老闆報告：漏客、hot leads、等付款、未回覆", "讓老闆每天看到工具價值"),
        ("MVP 4", "美容院 booking 或 IG shop order template", "先做一個垂直場景，打出案例"),
        ("MVP 5", "簡單 dashboard + staff ownership + follow-up reminder", "由 AI chatbot 變成生意跟進系統"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    widths = [0.85, 3.1, 2.8]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, fill=LIGHT_GREEN)
        table.rows[0].cells[idx].width = Inches(widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, bold=(idx == 0))
            cells[idx].width = Inches(widths[idx])


def add_callout(doc, label, text, fill=LIGHT_BLUE, color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    set_run_font(r, size=10.5, bold=True, color=color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.12
    r2 = p2.add_run(text)
    set_run_font(r2, size=10.5, color=TITLE_COLOR)
    doc.add_paragraph()


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Direction summary for HK SME DM Sales Cockpit")
    set_run_font(run, size=8.5, color=MUTED)


def build():
    doc = Document()
    style_doc(doc)
    add_footer(doc)

    title = doc.add_paragraph()
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("我們的方向：不要做普通 Chatbot")
    set_run_font(r, size=24, bold=True, color=TITLE_COLOR)
    add_para(
        doc,
        "建議定位：香港小店用得起的 IG/WhatsApp 漏客偵測 + DM 成交跟進系統",
        size=13,
        bold=True,
        color=BLUE,
        after=12,
    )
    add_callout(
        doc,
        "一句話",
        "我們不賣「AI 自動回覆」。我們賣「幫小店搵返 IG/WhatsApp 入面漏咗嘅生意」。",
        fill=LIGHT_GREEN,
        color=GREEN,
    )

    add_heading(doc, "1. 核心判斷", 1)
    add_bullets(
        doc,
        [
            "Instagram 上已經有大量 chatbot。只講 24/7、AI、WhatsApp/IG、CRM、broadcast，會變成同質化產品。",
            "Omnichat 已經把大型品牌需要的 chat commerce、CDP、marketing automation、OMO sales 做得很完整。",
            "我們不應正面複製 Omnichat，而是切入它太重、太貴、太 enterprise 的市場空位。",
            "小店老闆真正痛點不是「想要 chatbot」，而是：有沒有漏客、誰最有機會成交、誰等付款、誰需要員工跟進。",
        ]
    )

    add_heading(doc, "2. Omnichat vs 我們", 1)
    add_comparison_table(doc)

    add_heading(doc, "3. 產品方向", 1)
    add_callout(
        doc,
        "產品名字方向",
        "HK SME DM Sales Cockpit：把 IG/WhatsApp 對話變成可追蹤的生意 pipeline。",
        fill=LIGHT_BLUE,
        color=BLUE,
    )
    add_bullets(
        doc,
        [
            "Unified DM Inbox Lite：先支援 IG / WhatsApp / Website 的其中 1-2 個，不追求一開始全渠道。",
            "AI Reply Draft：AI 先寫草稿，危險內容如退款、療程效果、付款確認、booking confirmation 交人處理。",
            "Lead Status：每個客自動分類為新查詢、高意向、等付款、等確認、已成交、投訴風險。",
            "Follow-up Reminder：客人問完價、問完 booking、問完付款後未回，就提醒員工或老闆。",
            "Daily Boss Report：每天列出漏客、hot leads、等付款、未回覆、常見問題。",
        ]
    )

    add_heading(doc, "4. 為什麼這個方向比較易打市場", 1)
    add_bullets(
        doc,
        [
            "比「AI chatbot」更尖：客人一聽就明白是在防止漏單，而不是又一個自動回覆工具。",
            "比 Omnichat 輕：不用先做 full CDP、loyalty、game、enterprise integration。",
            "比純 inbox 有價值：系統會幫老闆判斷 priority，而不是只把 message 集中在一起。",
            "比人工更可控：AI 只負責整理、草稿、分類，真正高風險內容交員工決定。",
        ]
    )

    add_heading(doc, "5. MVP Roadmap", 1)
    add_roadmap_table(doc)

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_heading(doc, "6. 第一個市場建議", 1)
    add_callout(
        doc,
        "先打的客戶",
        "每天有 20-100 個 IG/WhatsApp 查詢，但未大到願意買 Omnichat 的香港小店。",
        fill=LIGHT_AMBER,
        color=AMBER,
    )
    add_bullets(
        doc,
        [
            "首選：美容院 / 美甲 / 醫美小店。原因：booking、問價、優惠、投訴風險、療程承諾都集中在 DM。",
            "第二選：IG shop。原因：現貨、尺碼、順豐、自取、付款、出貨、退貨都靠 DM，容易漏單。",
            "第三選：教育中心 / 補習社。原因：試堂、課程查詢、家長跟進、退款和承諾風險高。",
        ]
    )

    add_heading(doc, "7. 建議定價", 1)
    add_bullets(
        doc,
        [
            "Starter：HK$299-499/月，AI 草稿、基本 inbox、每日漏客報告。",
            "Growth：HK$799-999/月，加入 lead pipeline、follow-up reminder、booking/order/payment status。",
            "Setup：HK$1,500-3,000 一次性，幫店舖整理 FAQ、服務、價錢、規則、風險字眼。",
        ]
    )

    add_heading(doc, "8. 我們的 Instagram 內容方向", 1)
    add_bullets(
        doc,
        [
            "不要拍「我們有 AI chatbot」。要拍「你今日漏咗幾多 DM 生意？」",
            "內容例子：美容院最易漏單的 5 種 DM、IG shop 客人問完有冇貨之後點追、每日收工前應該看的漏客清單。",
            "Demo 畫面重點：hot leads、等付款、未回覆、投訴風險、今日建議先跟進的 5 個客。",
        ]
    )

    add_heading(doc, "9. 最終方向", 1)
    add_callout(
        doc,
        "Recommendation",
        "不要做 Omnichat clone。做 Omnichat 不想服務的小店市場：更輕、更平、更快上手，主打「不漏客」和「DM 轉成交」。",
        fill=LIGHT_GREEN,
        color=GREEN,
    )
    add_para(
        doc,
        "參考資料：Omnichat 官方 pricing page；Omnichat guest sharing slides；前期香港 SME / consumer messaging research。",
        size=8.8,
        color=MUTED,
        after=0,
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
